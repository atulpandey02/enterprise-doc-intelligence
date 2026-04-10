import os
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_from_pdf(file_path: str) -> Optional[str]:
    try:
        import fitz
        doc   = fitz.open(file_path)
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()
        full_text = "\n\n".join(pages)
        logger.info(f"PDF extracted: {len(pages)} pages, {len(full_text)} chars")
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return None


def extract_from_docx(file_path: str) -> Optional[str]:
    try:
        from docx import Document
        doc        = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n\n".join(paragraphs)
        logger.info(f"DOCX extracted: {len(paragraphs)} paragraphs")
        return full_text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return None


def extract_from_txt(file_path: str) -> Optional[str]:
    try:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        return None


def process_document(file_path: str) -> Optional[dict]:
    path      = Path(file_path)
    file_type = path.suffix.lower().lstrip(".")
    SUPPORTED = {
        "pdf":  extract_from_pdf,
        "docx": extract_from_docx,
        "txt":  extract_from_txt,
    }
    if file_type not in SUPPORTED:
        logger.warning(f"Unsupported file type: {file_type}")
        return None

    text = SUPPORTED[file_type](file_path)
    if not text or not text.strip():
        logger.warning(f"No text extracted from {file_path}")
        return None

    return {
        "file_path":  str(path.absolute()),
        "file_name":  path.name,
        "file_type":  file_type,
        "text":       text,
        "char_count": len(text),
    }
